# -*- coding: utf-8 -*-
"""
生成《校园二手交易软件——需求验证问卷》Word 文档。
运行: /d/python314/python scripts/gen_questionnaire.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TITLE = "校园二手物品交易软件 · 需求验证问卷"
OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "docs", "需求验证问卷.docx")

BRAND = RGBColor(0x2E, 0x7D, 0x32)   # 绿色主色
GREY = RGBColor(0x66, 0x66, 0x66)


def set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BRAND
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    return p


def add_section(doc, num, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(f"第{num}部分  {text}")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BRAND
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY


def add_question(doc, q_no, q_text, options=None, multi=False, blank=False):
    """options: list[str]; multi: 多选; blank: 加一条横线填空"""
    p = doc.add_paragraph()
    tag = ""
    if options:
        tag = "【多选】" if multi else "【单选】"
    elif blank:
        tag = "【填空】"
    r = p.add_run(f"{q_no}. {tag}{q_text}")
    r.font.size = Pt(10.5)
    r.font.bold = True
    if options:
        box = "☐" if multi else "○"
        for opt in options:
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(0.8)
            run = op.add_run(f"{box} {opt}")
            run.font.size = Pt(10.5)
    if blank:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(0.8)
        bp.add_run("_______________________________________________").font.color.rgb = GREY


def main():
    doc = Document()
    set_base_font(doc)
    for s in doc.sections:
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    add_title(doc, TITLE)
    add_subtitle(doc, "感谢参与！本问卷匿名填写，约需 3–5 分钟，结果仅用于产品调研。")
    add_note(doc, "说明：本产品定位为「仅限本校在校学生」使用的二手交易平台，强调实名/学籍认证、"
                  "线下当面交易，并提供校园（含楼栋/寝室）导航路书。")

    # 第一部分：受访者背景
    add_section(doc, "一", "你的基本情况")
    add_question(doc, 1, "你目前的身份是？",
                 ["本科在校生", "研究生在校生", "刚毕业(1年内)", "教职工", "其他(非本校学生)"])
    add_question(doc, 2, "你的年级是？",
                 ["大一", "大二", "大三", "大四", "研究生", "其他"])
    add_question(doc, 3, "你所在的校区 / 学校是？", blank=True)
    add_question(doc, 4, "你平时住在？",
                 ["校内宿舍", "校外租房", "走读/家住本地"])

    # 第二部分：二手交易现状
    add_section(doc, "二", "你的二手交易习惯与痛点")
    add_question(doc, 5, "过去一年，你是否买卖过二手物品？",
                 ["经常(每月都有)", "偶尔(几个月一次)", "很少(一年一两次)", "从来没有"])
    add_question(doc, 6, "你通过哪些渠道买卖二手物品？",
                 ["闲鱼/转转等公开平台", "微信群/QQ群(校园二手群)", "表白墙/校园集市",
                  "熟人朋友之间", "学校跳蚤市场(线下)", "从没买卖过"], multi=True)
    add_question(doc, 7, "你最常买卖的二手物品类别是？",
                 ["教材/资料", "电子数码(手机/电脑/耳机)", "生活用品/家电", "自行车/电动车",
                  "美妆/服饰", "考研考证资料", "代金券/卡券", "其他"], multi=True)
    add_question(doc, 8, "在二手交易中，你遇到过哪些困扰？",
                 ["怕遇到骗子/收到假货", "对方是校外人，不方便当面交易", "不知道东西该卖多少钱",
                  "沟通麻烦、爽约多", "东西送到对方宿舍/取货麻烦", "平台鱼龙混杂、信息太乱",
                  "没遇到过困扰"], multi=True)
    add_question(doc, 9, "你更倾向于哪种交易方式？",
                 ["校内当面交易(一手交钱一手交货)", "线上支付+快递/代拿", "都可以"])

    # 第三部分：核心卖点验证——强认证
    add_section(doc, "三", "关于「仅限本校学生」与实名认证")
    add_question(doc, 10, "如果一个平台保证「所有用户都是经过认证的本校在校学生」，你的购买/交易意愿会？",
                 ["明显提高，更放心", "有一点提高", "没什么影响", "反而担心隐私"])
    add_question(doc, 11, "为了通过学生认证，你愿意接受哪种方式？",
                 ["用学校邮箱(@xxx.edu.cn)收验证码", "上传学生证照片审核",
                  "上传学信网学籍在线验证报告", "刷校园卡/学号+姓名核验", "都不愿意"], multi=True)
    add_question(doc, 12, "你对上传学生证 / 学籍信息做认证的顾虑是？",
                 ["担心个人信息泄露", "担心照片被滥用", "嫌麻烦、流程太长", "没有顾虑，可以接受"], multi=True)
    add_question(doc, 13, "你认为「只有本校学生能用」这个限制，对你来说是？",
                 ["优点，圈子干净更安全", "无所谓", "缺点，可选的人太少、东西太少"])

    # 第四部分：核心卖点验证——校园导航路书
    add_section(doc, "四", "关于「校园导航 / 寝室路书」")
    add_question(doc, 14, "当面交易时，你是否遇到过「找不到对方宿舍楼/教学楼具体位置」的情况？",
                 ["经常遇到", "偶尔遇到", "很少", "从没有(校园很熟)"])
    add_question(doc, 15, "如果 App 里有「从校门口→X 号楼→X 层→XXX 寝室」的图文导航路书，你觉得？",
                 ["很有用，愿意用", "有点用", "用处不大", "完全用不到"])
    add_question(doc, 16, "如果「上传优质校园导航路书可获得代金券奖励」，你愿意参与贡献吗？",
                 ["愿意，为了奖励会认真做", "看心情，顺手做", "不愿意，太麻烦"])
    add_question(doc, 17, "你更希望路书包含哪些内容？",
                 ["分步文字指引", "实拍照片", "手绘/标注路线图", "楼栋平面示意图",
                  "语音/视频讲解"], multi=True)

    # 第五部分：使用与付费意愿
    add_section(doc, "五", "使用与付费意愿")
    add_question(doc, 18, "如果这样一款校园二手小程序上线，你会用吗？",
                 ["一定会用", "可能会用", "看情况", "不会用"])
    add_question(doc, 19, "你更希望它先做成哪种形态？",
                 ["微信小程序(扫码即用)", "安卓App", "两个都要", "无所谓"])
    add_question(doc, 20, "以下增值服务，你可能愿意付费的是？(仅撮合、当面交易，平台不碰货款)",
                 ["商品置顶/加急曝光", "认证摊主/信用标识", "代金券商城", "会员(更多发布额度)",
                  "都不愿意付费"], multi=True)
    add_question(doc, 21, "你希望这款产品叫什么名字 / 有什么建议？(选填)", blank=True)
    add_question(doc, 22, "如果愿意后续参与内测或访谈，请留下联系方式(选填，仅用于邀请)", blank=True)

    # 结尾
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("—— 问卷到此结束，感谢你的宝贵意见！ ——")
    er.font.size = Pt(10.5)
    er.font.color.rgb = BRAND
    er.font.bold = True

    doc.save(OUT)
    print("SAVED:", OUT)


if __name__ == "__main__":
    main()
