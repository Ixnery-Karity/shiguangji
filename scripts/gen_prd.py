# -*- coding: utf-8 -*-
"""
生成《校园二手交易小程序 MVP 产品需求文档(PRD)》Word。
运行: /d/python314/python scripts/gen_prd.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "docs", "MVP产品需求文档PRD.docx")
BRAND = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)


def cn(run, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"
    st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = BRAND
    cn(r)
    # 底部横线
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"),
                              {qn("w:val"): "single", qn("w:sz"): "6",
                               qn("w:space"): "2", qn("w:color"): "2E7D32"})
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = DARK
    cn(r)
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    cn(r)
    return p


def bullet(doc, text, indent=0.6):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run("• " + text)
    r.font.size = Pt(10.5)
    cn(r)
    return p


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cn(run)
        # 表头背景色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"),
                               {qn("w:val"): "clear", qn("w:fill"): "2E7D32"})
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            cn(run)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("校园二手物品交易平台")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = BRAND; cn(r)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("MVP 产品需求文档 (PRD)")
    r2.font.size = Pt(18); r2.font.color.rgb = DARK; cn(r2)
    doc.add_paragraph()
    for line in ["版本：v0.1.0", "阶段：MVP（微信小程序先行）",
                 "定位：仅限本校在校学生 · 强认证 · 仅撮合线下当面交易",
                 "日期：2026-07-12"]:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(line)
        rl.font.size = Pt(11); rl.font.color.rgb = GREY; cn(rl)
    doc.add_page_break()


def main():
    doc = Document()
    set_base(doc)
    for s in doc.sections:
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    title_page(doc)

    # 1 文档说明
    h1(doc, "1. 文档说明与产品概述")
    h2(doc, "1.1 产品一句话定位")
    body(doc, "一个仅限本校在校学生使用的二手物品交易微信小程序：通过严格的学生身份认证保证圈子纯净与交易安全，"
              "平台只负责信息撮合与沟通，交易在校内当面完成（一手交钱一手交货），平台不经手货款。")
    h2(doc, "1.2 目标与非目标（MVP 边界）")
    add_table(doc, ["本期要做（In Scope）", "本期不做（Out of Scope）"], [
        ["学生认证（邮箱 + 学生证/学籍核验）", "线上支付、资金担保、平台抽佣"],
        ["发布 / 浏览 / 搜索二手商品", "快递物流、代拿代取"],
        ["站内即时聊天（IM）", "室内实时定位导航（蓝牙/UWB）"],
        ["当面交易「已完成」标记 + 评价", "安卓 / iOS 原生 App（放到二期）"],
        ["校园导航路书（图文 UGC，二期灰度）", "复杂的信用分体系、代金券商城"],
    ], widths=[8.0, 8.0])
    h2(doc, "1.3 目标用户")
    bullet(doc, "卖家：有闲置物品要处理的在校学生（毕业生、换季清仓、数码升级等）。")
    bullet(doc, "买家：想低价购入教材、生活用品、数码产品的在校学生。")
    bullet(doc, "贡献者：熟悉校园、愿意上传导航路书赚代金券的活跃学生。")

    # 2 名词与角色
    h1(doc, "2. 用户角色与权限")
    add_table(doc, ["角色", "说明", "关键权限"], [
        ["游客", "未登录/未认证", "仅可浏览商品列表与详情，不能发布/联系卖家"],
        ["认证学生", "通过学生认证的本校学生", "发布、聊天、交易、评价、上传路书"],
        ["运营/审核员", "平台后台人员", "认证材料审核、商品/内容审核、举报处理、代金券发放"],
        ["超级管理员", "平台负责人", "全部后台权限 + 数据看板 + 配置管理"],
    ], widths=[3.0, 6.0, 8.0])

    # 3 功能清单
    h1(doc, "3. 功能清单（含优先级）")
    body(doc, "优先级说明：P0=MVP 必做，上线即有；P1=MVP 期望有，时间紧可延后；P2=二期。", 0)
    add_table(doc, ["模块", "功能点", "优先级", "说明"], [
        ["认证", "微信授权登录", "P0", "获取 openid，建立账号"],
        ["认证", "校园邮箱验证码认证", "P0", "@xxx.edu.cn 收码，一级门槛"],
        ["认证", "学生证 OCR + 人工审核", "P0", "上传照片，后台复核，二级门槛"],
        ["认证", "学信网学籍报告核验", "P1", "上传在线验证报告，最高可信度"],
        ["认证", "认证状态与徽章展示", "P0", "已认证学生显示校徽标识"],
        ["商品", "发布商品（图文+分类+价格）", "P0", "最多9图，选楼栋位置"],
        ["商品", "商品列表 / 分类 / 瀑布流", "P0", "首页信息流"],
        ["商品", "搜索 + 筛选（分类/价格/成色）", "P0", "关键词与条件筛选"],
        ["商品", "商品详情页", "P0", "图文、卖家、位置、联系按钮"],
        ["商品", "收藏 / 想要", "P1", "加入我的收藏"],
        ["商品", "编辑 / 下架 / 标记已售", "P0", "卖家管理自己的商品"],
        ["交易", "站内即时聊天 IM", "P0", "买卖双方沟通，第三方IM或WebSocket"],
        ["交易", "发起交易 / 约定当面交易", "P0", "生成交易记录，约时间地点"],
        ["交易", "交易完成双向确认", "P0", "双方点「已完成」闭环"],
        ["交易", "交易后互相评价", "P1", "好/中/差 + 文字"],
        ["安全", "举报 / 拉黑", "P0", "举报商品、用户，拉黑"],
        ["安全", "敏感词 / 违禁品拦截", "P0", "发布时校验，违禁品清单"],
        ["地图", "校园导航路书浏览", "P2", "图文路书列表与详情"],
        ["地图", "上传路书 + 代金券奖励", "P2", "UGC 贡献激励"],
        ["个人", "个人主页 / 我的发布 / 我的交易", "P0", "个人中心"],
        ["个人", "消息通知中心", "P1", "系统与交易通知"],
        ["后台", "认证审核工作台", "P0", "通过/驳回，留痕"],
        ["后台", "商品 / 内容审核", "P0", "下架违规内容"],
        ["后台", "举报处理 / 用户管理", "P0", "封禁、警告"],
        ["后台", "数据看板", "P1", "DAU、发布量、成交量"],
    ], widths=[2.2, 6.4, 1.8, 6.0])

    # 4 关键流程
    h1(doc, "4. 核心业务流程")
    h2(doc, "4.1 学生认证流程（强认证）")
    for i, step in enumerate([
        "微信授权登录 → 创建基础账号（此时为「游客/未认证」）。",
        "第一级：绑定校园邮箱 @xxx.edu.cn，收取 6 位验证码并校验 → 标记「邮箱已验证」。",
        "第二级：上传学生证照片，OCR 自动提取姓名/学号/学校 → 提交后台人工审核。",
        "（P1 可选）上传学信网学籍在线验证报告，后台核对验证码真伪。",
        "后台审核通过 → 账号升级为「认证学生」，主页显示校徽徽章，解锁发布/交易权限。",
        "审核驳回 → 通知用户原因，可重新提交。敏感材料审核后按合规要求脱敏/限期删除。",
    ], 1):
        body(doc, f"步骤{i}：{step}", 0.4)
    body(doc, "认证门槛设计原则：邮箱=准入，学生证/学籍=可信身份。两级叠加，兼顾转化率与安全。", 0.4)

    h2(doc, "4.2 发布商品流程")
    for i, step in enumerate([
        "点击「发布」→ 校验是否已认证（未认证则引导去认证）。",
        "上传 1–9 张实拍图 → 填写标题、描述、分类、成色、价格。",
        "选择交易位置（校区 / 楼栋，用于买家就近判断）。",
        "提交 → 敏感词与违禁品校验 → 进入列表（可配置先审后发或先发后审）。",
    ], 1):
        body(doc, f"步骤{i}：{step}", 0.4)

    h2(doc, "4.3 撮合与线下交易流程（仅撮合，不碰货款）")
    for i, step in enumerate([
        "买家在详情页点「聊一聊」→ 进入 IM 与卖家沟通。",
        "双方约定当面交易的时间、地点（可引用校园路书辅助找到位置）。",
        "任意一方点「发起交易」→ 生成交易记录（状态：待完成）。",
        "线下当面一手交钱一手交货（平台不经手资金）。",
        "双方各自点「确认已完成」→ 双向确认后交易闭环，商品自动标记「已售」。",
        "交易完成后可互相评价，沉淀信用。",
    ], 1):
        body(doc, f"步骤{i}：{step}", 0.4)
    body(doc, "风险提示：平台在交易页显著位置提示「请在校内公共区域当面验货交易，谨防诈骗」，"
              "并保留聊天与交易记录用于举报取证。", 0.4)

    # 5 页面结构
    h1(doc, "5. 页面结构与信息架构")
    h2(doc, "5.1 小程序页面清单")
    add_table(doc, ["页面", "路径(建议)", "说明"], [
        ["首页/信息流", "pages/home", "商品瀑布流 + 搜索入口 + 分类"],
        ["搜索/筛选", "pages/search", "关键词、分类、价格、成色筛选"],
        ["商品详情", "pages/item/detail", "图文、卖家、位置、联系/收藏"],
        ["发布商品", "pages/item/publish", "表单 + 上传图片"],
        ["聊天列表", "pages/chat/list", "会话列表"],
        ["聊天会话", "pages/chat/room", "IM 对话 + 发起交易"],
        ["认证中心", "pages/auth/verify", "邮箱/学生证/学籍认证"],
        ["个人中心", "pages/mine", "我的发布/收藏/交易/设置"],
        ["我的交易", "pages/trade/list", "进行中/已完成交易"],
        ["导航路书(二期)", "pages/guide", "路书列表与详情"],
    ], widths=[3.4, 5.2, 7.0])

    h2(doc, "5.2 底部导航（TabBar）")
    bullet(doc, "首页　|　发布　|　消息　|　我的（4 个 Tab，发布居中突出）")

    # 6 数据模型
    h1(doc, "6. 数据模型概览（核心实体）")
    add_table(doc, ["实体", "关键字段", "说明"], [
        ["User 用户", "id, openid, nickname, avatar, school, auth_level, credit", "auth_level: 游客/邮箱认证/学生认证"],
        ["Auth 认证记录", "id, user_id, type, material_url, status, reviewer, review_time", "type: 邮箱/学生证/学籍"],
        ["Item 商品", "id, seller_id, title, desc, price, category, condition, images, location, status", "status: 在售/已售/下架"],
        ["Chat 会话", "id, item_id, buyer_id, seller_id, last_msg, updated_at", "一商品一买家一会话"],
        ["Message 消息", "id, chat_id, sender_id, content, type, created_at", "type: 文本/图片/系统"],
        ["Trade 交易", "id, item_id, buyer_id, seller_id, status, buyer_done, seller_done", "双向确认完成"],
        ["Review 评价", "id, trade_id, from_id, to_id, score, content", "交易后互评"],
        ["Guide 路书", "id, author_id, building, floor, room, steps, images, reward", "二期，UGC 导航"],
        ["Report 举报", "id, reporter_id, target_type, target_id, reason, status", "举报处理"],
    ], widths=[3.0, 7.2, 5.4])

    # 7 非功能需求
    h1(doc, "7. 非功能性需求")
    bullet(doc, "合规：遵守《个人信息保护法》，认证材料（证件照、学籍）属敏感信息，最小化收集、加密存储、"
                "审核后脱敏/限期删除，需明确《隐私政策》与《用户协议》。")
    bullet(doc, "安全：接口鉴权（JWT）、上传图片鉴黄、敏感词过滤、违禁品清单（管制刀具/药品等禁止交易）。")
    bullet(doc, "性能：首页信息流首屏 < 2s，图片走 CDN + 懒加载；列表分页/瀑布流加载。")
    bullet(doc, "可用性：小程序需通过微信审核，涉及交易的类目需符合微信平台规则。")
    bullet(doc, "可扩展：后端与小程序解耦，二期安卓 App 复用同一套后端 API。")

    # 8 技术选型
    h1(doc, "8. 技术选型建议")
    add_table(doc, ["层", "选型", "理由"], [
        ["小程序端", "微信原生 或 uni-app", "uni-app 便于二期复用到 App/H5"],
        ["后端", "Spring Boot / NestJS", "团队熟悉优先；RESTful API"],
        ["数据库", "PostgreSQL + Redis", "关系数据 + 缓存/会话"],
        ["即时通讯", "第三方IM(融云/环信) 或 WebSocket", "自研IM不划算，建议接第三方"],
        ["对象存储", "阿里云 OSS / 腾讯云 COS", "存商品图、认证材料、路书图"],
        ["地图", "高德地图小程序 SDK", "校园室外定位 + 路书底图"],
        ["OCR/核验", "腾讯云/阿里云 OCR + 学籍核验API", "学生证识别、学籍核验(需企业资质)"],
        ["后台管理", "Vue/React + Ant Design", "认证审核、内容管理"],
    ], widths=[3.0, 6.0, 6.6])

    # 9 里程碑
    h1(doc, "9. 开发里程碑（建议排期）")
    add_table(doc, ["阶段", "周期", "交付内容"], [
        ["M0 需求验证", "1–2 周", "问卷+访谈，确认需求与目标校"],
        ["M1 认证+账号", "2 周", "微信登录、邮箱认证、学生证审核、后台审核台"],
        ["M2 商品闭环", "3 周", "发布/列表/搜索/详情/收藏"],
        ["M3 交易闭环", "2 周", "IM、发起交易、双向确认、评价"],
        ["M4 安全+上线", "1–2 周", "举报、敏感词、审核、隐私协议、提审上线"],
        ["M5 二期", "—", "导航路书+代金券、安卓 App"],
    ], widths=[3.0, 2.6, 10.0])

    # 10 风险
    h1(doc, "10. 主要风险与对策")
    add_table(doc, ["风险", "对策"], [
        ["学信网无个人接口，认证难", "分级认证：邮箱+学生证起步，企业主体后接第三方学籍核验API"],
        ["个人敏感信息合规风险", "最小化收集、加密、审核后脱敏删除、完善隐私政策"],
        ["冷启动没人用（先有鸡还是蛋）", "单校聚焦、毕业季切入、地推+表白墙+社群、路书代金券引流"],
        ["线下交易纠纷/诈骗", "全程留痕、举报取证、公共区域交易提示、信用评价"],
        ["微信平台审核类目限制", "提前研究二手交易类目规则，规避违禁品与资金相关功能"],
    ], widths=[6.0, 10.0])

    doc.save(OUT)
    print("SAVED:", OUT)


if __name__ == "__main__":
    main()
